import re
import unicodedata
from pathlib import Path

import geopandas as gpd
import pandas as pd
from docx import Document


VALIDATION_REPORT_PATH = Path("outputs") / "gpkg_validation_report.csv"
RAW_DATA_DIR = Path(r"C:\Users\saiga\Desktop\PhD\CSMD_APP\data_raw")
TABLE1_PATH = Path("reference") / "Veeravalli_ED_Table1.docx"
CITY_OUTPUT_PATH = Path("outputs") / "city_size_5class_validation.csv"
COUNTRY_OUTPUT_PATH = Path("outputs") / "country_table1_crosscheck.csv"

REQUIRED_GPKG_COLUMNS = [
    "ID_HDC_G0",
    "UC_NM_MN",
    "CTR_MN_NM",
    "POP25_U",
    "POP_SEG",
    "rf_label",
]
REQUIRED_TABLE_COLUMNS = ["ISO3", "Country", "Total_Pop_M", "Deprived_Pop_M"]
POPULATION_TOLERANCE_M = 0.01

CITY_SIZE_LABELS = {
    1: "Small (<500,000 residents)",
    2: "Medium (500,000–<1 million residents)",
    3: "Large (1–<5 million residents)",
    4: "Very large (5–<10 million residents)",
    5: "Megacity (≥10 million residents)",
}
CITY_SIZE_NAMES = {
    1: "Small",
    2: "Medium",
    3: "Large",
    4: "Very large",
    5: "Megacity",
}


def valid_for_viewer_mask(series: pd.Series) -> pd.Series:
    return (
        series.fillna(False)
        .astype(str)
        .str.strip()
        .str.lower()
        .isin({"true", "1", "yes"})
    )


def city_size_code(population: pd.Series) -> pd.Series:
    return pd.cut(
        population,
        bins=[float("-inf"), 500_000, 1_000_000, 5_000_000, 10_000_000, float("inf")],
        labels=[1, 2, 3, 4, 5],
        right=False,
    ).astype("Int64")


def normalized_numeric_values(series: pd.Series) -> set[float | None]:
    return {
        None if pd.isna(value) else float(value)
        for value in pd.unique(series)
    }


def normalized_integer_values(series: pd.Series) -> set[int | None]:
    return {
        None if pd.isna(value) else int(value)
        for value in pd.unique(series)
    }


def normalized_text_values(series: pd.Series) -> set[str | None]:
    return {
        None if pd.isna(value) else str(value).strip()
        for value in pd.unique(series)
    }


def representative_value(values: set):
    non_missing = sorted(value for value in values if value is not None)
    if non_missing:
        return non_missing[0]
    return pd.NA


def values_are_consistent(values: set) -> bool:
    return len(values) == 1 and None not in values


def normalize_country_name(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", str(value))
    ascii_value = normalized.encode("ascii", "ignore").decode("ascii")
    return re.sub(r"[^a-z0-9]+", " ", ascii_value.lower()).strip()


def read_valid_files() -> pd.DataFrame:
    if not VALIDATION_REPORT_PATH.exists():
        raise FileNotFoundError(
            f"Validation report not found: {VALIDATION_REPORT_PATH}. "
            "Run preprocessing/01_validate_gpkgs.py first."
        )

    report = pd.read_csv(VALIDATION_REPORT_PATH)
    required_columns = {"file", "first_layer", "valid_for_viewer"}
    missing_columns = sorted(required_columns - set(report.columns))
    if missing_columns:
        raise ValueError(
            "Validation report is missing required columns: "
            + ", ".join(missing_columns)
        )

    valid_files = report[valid_for_viewer_mask(report["valid_for_viewer"])].copy()
    if valid_files.empty:
        raise ValueError("No GeoPackages are marked valid_for_viewer in the report.")
    return valid_files


def collect_gpkg_statistics(valid_files: pd.DataFrame) -> tuple[dict, dict]:
    cities = {}
    countries = {}

    for report_row in valid_files.itertuples(index=False):
        gpkg_path = RAW_DATA_DIR / Path(str(report_row.file)).name
        if not gpkg_path.exists():
            raise FileNotFoundError(f"GeoPackage not found: {gpkg_path}")

        layer = str(report_row.first_layer).strip()
        layer = layer if layer and layer.lower() != "nan" else None
        df = gpd.read_file(
            gpkg_path,
            layer=layer,
            columns=REQUIRED_GPKG_COLUMNS,
            ignore_geometry=True,
            engine="pyogrio",
        )
        missing_columns = [
            column for column in REQUIRED_GPKG_COLUMNS if column not in df.columns
        ]
        if missing_columns:
            raise ValueError(
                f"{gpkg_path.name} is missing required columns: "
                f"{', '.join(missing_columns)}"
            )
        if df["ID_HDC_G0"].isna().any():
            raise ValueError(f"{gpkg_path.name} contains missing ID_HDC_G0 values.")

        working = df[["ID_HDC_G0", "UC_NM_MN", "CTR_MN_NM"]].copy()
        working["POP25_U"] = pd.to_numeric(df["POP25_U"], errors="coerce")
        working["POP_SEG"] = pd.to_numeric(df["POP_SEG"], errors="coerce")
        working["rf_label"] = pd.to_numeric(df["rf_label"], errors="coerce")
        working["city_size_code"] = city_size_code(working["POP25_U"])

        for city_id, group in working.groupby("ID_HDC_G0", sort=False):
            record = cities.setdefault(
                float(city_id),
                {
                    "ID_HDC_G0": float(city_id),
                    "city_values": set(),
                    "country_values": set(),
                    "population_values": set(),
                    "city_size_codes": set(),
                    "number_of_segments": 0,
                },
            )
            record["city_values"].update(
                normalized_text_values(group["UC_NM_MN"])
            )
            record["country_values"].update(
                normalized_text_values(group["CTR_MN_NM"])
            )
            record["population_values"].update(
                normalized_numeric_values(group["POP25_U"])
            )
            record["city_size_codes"].update(
                normalized_integer_values(group["city_size_code"])
            )
            record["number_of_segments"] += len(group)

        for country, group in working.groupby("CTR_MN_NM", dropna=False, sort=False):
            country_name = None if pd.isna(country) else str(country).strip()
            country_key = normalize_country_name(country_name or "")
            record = countries.setdefault(
                country_key,
                {
                    "gpkg_country": country_name,
                    "total_population": 0.0,
                    "deprived_population": 0.0,
                },
            )
            record["total_population"] += float(group["POP_SEG"].sum())
            record["deprived_population"] += float(
                group.loc[group["rf_label"] == 1, "POP_SEG"].sum()
            )

    return cities, countries


def create_city_validation(cities: dict) -> pd.DataFrame:
    records = []

    for city_record in cities.values():
        city_values = city_record["city_values"]
        country_values = city_record["country_values"]
        population_values = city_record["population_values"]
        size_codes = city_record["city_size_codes"]
        code = representative_value(size_codes)

        records.append(
            {
                "ID_HDC_G0": city_record["ID_HDC_G0"],
                "country": representative_value(country_values),
                "city": representative_value(city_values),
                "POP25_U": representative_value(population_values),
                "city_size_code": code,
                "city_size_label": (
                    CITY_SIZE_LABELS.get(int(code), "") if not pd.isna(code) else ""
                ),
                "number_of_segments": int(city_record["number_of_segments"]),
                "inconsistent_city_names": not values_are_consistent(city_values),
                "inconsistent_country_names": not values_are_consistent(country_values),
                "inconsistent_population_values": not values_are_consistent(
                    population_values
                ),
                "inconsistent_city_size_codes": not values_are_consistent(size_codes),
            }
        )

    result = pd.DataFrame(records)
    id_values = pd.to_numeric(result["ID_HDC_G0"], errors="coerce")
    if ((id_values.dropna() % 1) == 0).all():
        result["ID_HDC_G0"] = id_values.astype("Int64")
    else:
        result["ID_HDC_G0"] = id_values
    result["POP25_U"] = pd.to_numeric(result["POP25_U"], errors="coerce")
    result["city_size_code"] = pd.to_numeric(
        result["city_size_code"], errors="coerce"
    ).astype("Int64")

    return result.sort_values(
        ["country", "city", "ID_HDC_G0"],
        key=lambda column: column.astype(str).str.lower(),
        kind="stable",
    ).reset_index(drop=True)


def read_table1() -> pd.DataFrame:
    if not TABLE1_PATH.exists():
        raise FileNotFoundError(f"Table 1 document not found: {TABLE1_PATH}")

    document = Document(TABLE1_PATH)
    if not document.tables:
        raise ValueError(f"No tables found in: {TABLE1_PATH}")

    table = document.tables[0]
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    missing_columns = [
        column for column in REQUIRED_TABLE_COLUMNS if column not in headers
    ]
    if missing_columns:
        raise ValueError(
            "Table 1 is missing required columns: " + ", ".join(missing_columns)
        )

    rows = []
    for table_row in table.rows[1:]:
        values = [cell.text.strip() for cell in table_row.cells]
        row = dict(zip(headers, values))
        if not row.get("ISO3") and not row.get("Country"):
            continue
        rows.append({column: row.get(column, "") for column in REQUIRED_TABLE_COLUMNS})

    result = pd.DataFrame(rows)
    result["ISO3"] = result["ISO3"].astype(str).str.strip()
    result["Country"] = result["Country"].astype(str).str.strip()
    for column in ["Total_Pop_M", "Deprived_Pop_M"]:
        result[column] = pd.to_numeric(
            result[column].astype(str).str.replace(",", "", regex=False),
            errors="coerce",
        )
    result["country_key"] = result["Country"].map(normalize_country_name)

    duplicate_keys = result.loc[
        result["country_key"].duplicated(keep=False), "Country"
    ].tolist()
    if duplicate_keys:
        raise ValueError(
            "Table 1 contains duplicate normalized country names: "
            + ", ".join(duplicate_keys)
        )
    return result


def create_country_crosscheck(table1: pd.DataFrame, countries: dict) -> pd.DataFrame:
    calculated = pd.DataFrame(
        {
            "country_key": country_key,
            "gpkg_country": values["gpkg_country"],
            "calculated_total_pop_m": values["total_population"] / 1_000_000,
            "calculated_deprived_pop_m": values["deprived_population"] / 1_000_000,
        }
        for country_key, values in countries.items()
    )

    result = table1.merge(calculated, on="country_key", how="outer", validate="one_to_one")
    result["matched_to_table1"] = result["Country"].notna() & result[
        "gpkg_country"
    ].notna()
    result["total_pop_difference_m"] = (
        result["calculated_total_pop_m"] - result["Total_Pop_M"]
    )
    result["deprived_pop_difference_m"] = (
        result["calculated_deprived_pop_m"] - result["Deprived_Pop_M"]
    )
    result["total_pop_matches"] = result["matched_to_table1"] & (
        result["total_pop_difference_m"].abs()
        <= POPULATION_TOLERANCE_M + 1e-12
    )
    result["deprived_pop_matches"] = result["matched_to_table1"] & (
        result["deprived_pop_difference_m"].abs()
        <= POPULATION_TOLERANCE_M + 1e-12
    )

    numeric_columns = [
        "calculated_total_pop_m",
        "calculated_deprived_pop_m",
        "total_pop_difference_m",
        "deprived_pop_difference_m",
    ]
    result[numeric_columns] = result[numeric_columns].round(6)

    output_columns = [
        "ISO3",
        "Country",
        "gpkg_country",
        "Total_Pop_M",
        "Deprived_Pop_M",
        "calculated_total_pop_m",
        "calculated_deprived_pop_m",
        "total_pop_difference_m",
        "deprived_pop_difference_m",
        "matched_to_table1",
        "total_pop_matches",
        "deprived_pop_matches",
    ]
    return result[output_columns].sort_values(
        ["Country", "gpkg_country"],
        key=lambda column: column.astype(str).str.lower(),
        kind="stable",
    ).reset_index(drop=True)


def print_summary(city_result: pd.DataFrame, country_result: pd.DataFrame) -> None:
    print(f"Cities: {len(city_result)}")
    for code, name in CITY_SIZE_NAMES.items():
        count = int((city_result["city_size_code"] == code).sum())
        print(f"City size {code} - {name}: {count}")

    inconsistent_population = int(
        city_result["inconsistent_population_values"].sum()
    )
    print(f"Cities with inconsistent POP25_U values: {inconsistent_population}")

    matched = country_result["matched_to_table1"]
    unmatched = country_result[~matched]
    failed = country_result[
        matched
        & (~country_result["total_pop_matches"] | ~country_result["deprived_pop_matches"])
    ]
    print(f"Countries matched to Table 1: {int(matched.sum())}")
    print(f"Unmatched countries: {len(unmatched)}")
    if not unmatched.empty:
        unmatched_names = unmatched["Country"].fillna(unmatched["gpkg_country"])
        print("Unmatched country names: " + ", ".join(unmatched_names.astype(str)))
    print(f"Countries failing either population comparison: {len(failed)}")
    if not failed.empty:
        failed_names = failed["Country"].fillna(failed["gpkg_country"])
        print("Failing country names: " + ", ".join(failed_names.astype(str)))

    matched_rows = country_result[matched]
    max_total_difference = matched_rows["total_pop_difference_m"].abs().max()
    max_deprived_difference = matched_rows[
        "deprived_pop_difference_m"
    ].abs().max()
    print(f"Maximum absolute total population difference (M): {max_total_difference:.6f}")
    print(
        "Maximum absolute deprived population difference (M): "
        f"{max_deprived_difference:.6f}"
    )


def main() -> None:
    if not RAW_DATA_DIR.exists():
        raise FileNotFoundError(f"Raw data directory not found: {RAW_DATA_DIR}")

    valid_files = read_valid_files()
    cities, countries = collect_gpkg_statistics(valid_files)
    city_result = create_city_validation(cities)
    table1 = read_table1()
    country_result = create_country_crosscheck(table1, countries)

    CITY_OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    city_result.to_csv(CITY_OUTPUT_PATH, index=False)
    country_result.to_csv(COUNTRY_OUTPUT_PATH, index=False)

    print_summary(city_result, country_result)
    print(f"City validation output: {CITY_OUTPUT_PATH}")
    print(f"Country cross-check output: {COUNTRY_OUTPUT_PATH}")


if __name__ == "__main__":
    main()
